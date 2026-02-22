from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
title = doc.add_heading('🛡️ Manual Tehnic AXON CORE OS v6.17', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Secțiunea Hardware
doc.add_heading('1. Configurație Hardware', level=1)
doc.add_paragraph('Sistem de Operare: Debian GNU/Linux 12 (Bookworm)', style='List Bullet')
doc.add_paragraph('Infrastructură: Google Cloud Compute Engine', style='List Bullet')
doc.add_paragraph('Virtualizare: Docker Engine & Compose v2', style='List Bullet')

# Secțiunea Software
doc.add_heading('2. Software Stack', level=1)
doc.add_paragraph('Core: Python 3.11-slim', style='List Bullet')
doc.add_paragraph('Frontend: Streamlit 1.30+', style='List Bullet')
doc.add_paragraph('AI Engine: Gemini 2.5 Flash (Multimodal)', style='List Bullet')

# FrameLang
doc.add_heading('3. Logica FrameLang', level=1)
doc.add_paragraph('FrameLang orchestrează identitatea celor 11 departamente și permite analiza multimodală a fișierelor Word, PDF și Imagini.')

doc.save('Documentatie_Elaborata_AXON_v6.17.docx')
